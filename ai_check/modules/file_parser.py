# -*- coding: utf-8 -*-
import os
import logging
from typing import Dict, Any, Optional
from pathlib import Path

class FileParser:
    """文件解析类，负责解析Word和PDF文件"""
    
    def __init__(self, temp_dir: str = "temp"):
        """初始化文件解析器
        
        Args:
            temp_dir: 临时文件目录
        """
        self.temp_dir = temp_dir
        self._ensure_temp_dir()
    
    def _ensure_temp_dir(self) -> None:
        """确保临时目录存在"""
        if not os.path.exists(self.temp_dir):
            os.makedirs(self.temp_dir)
            logging.info(f"创建临时目录: {self.temp_dir}")
    
    def parse_file(self, file_path: str) -> Dict[str, Any]:
        """解析文件内容
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析结果字典，包含文本内容和元数据
        """
        if not os.path.exists(file_path):
            raise FileNotFoundError(f"文件不存在: {file_path}")
        
        file_ext = Path(file_path).suffix.lower()
        
        if file_ext == ".docx":
            return self._parse_docx(file_path)
        elif file_ext == ".pdf":
            return self._parse_pdf(file_path)
        else:
            raise ValueError(f"不支持的文件格式: {file_ext}，仅支持.docx和.pdf格式")
    
    def _parse_docx(self, file_path: str) -> Dict[str, Any]:
        """解析Word文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析结果字典
        """
        try:
            import docx
            doc = docx.Document(file_path)
            
            # 提取文本内容，保留段落结构
            paragraphs = []
            for para in doc.paragraphs:
                if para.text.strip():
                    paragraphs.append(para.text)
            
            # 提取标题层级
            headings = []
            for para in doc.paragraphs:
                if para.style.name.startswith('Heading'):
                    level = int(para.style.name.replace('Heading', ''))
                    headings.append({
                        "level": level,
                        "text": para.text
                    })
            
            # 保存到临时文本文件
            temp_file_path = os.path.join(self.temp_dir, Path(file_path).stem + ".txt")
            with open(temp_file_path, 'w', encoding='utf-8') as f:
                f.write('\n\n'.join(paragraphs))
            
            return {
                "content": '\n\n'.join(paragraphs),
                "paragraphs": paragraphs,
                "headings": headings,
                "temp_file": temp_file_path,
                "file_type": "docx",
                "file_name": Path(file_path).name
            }
        except ImportError:
            logging.error("缺少python-docx库，请安装: pip install python-docx")
            raise
        except Exception as e:
            logging.error(f"解析Word文档失败: {str(e)}")
            raise ValueError(f"解析Word文档失败: {str(e)}")
    
    def _parse_pdf(self, file_path: str) -> Dict[str, Any]:
        """解析PDF文档
        
        Args:
            file_path: 文件路径
            
        Returns:
            解析结果字典
        """
        try:
            from PyPDF2 import PdfReader
            
            reader = PdfReader(file_path)
            text_content = ""
            paragraphs = []
            
            # 提取文本内容
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_content += page_text + "\n\n"
                    # 简单按换行符分割段落
                    page_paragraphs = [p.strip() for p in page_text.split('\n') if p.strip()]
                    paragraphs.extend(page_paragraphs)
            
            # 尝试使用pdfplumber获取更精确的文本（如果可用）
            try:
                import pdfplumber
                with pdfplumber.open(file_path) as pdf:
                    plumber_paragraphs = []
                    for page in pdf.pages:
                        page_text = page.extract_text()
                        if page_text:
                            page_paragraphs = [p.strip() for p in page_text.split('\n') if p.strip()]
                            plumber_paragraphs.extend(page_paragraphs)
                    
                    # 如果pdfplumber提取的文本更多，则使用它
                    if len(plumber_paragraphs) > len(paragraphs):
                        paragraphs = plumber_paragraphs
                        text_content = '\n\n'.join(plumber_paragraphs)
            except ImportError:
                logging.warning("未安装pdfplumber库，使用PyPDF2提取的文本")
            
            # 保存到临时文本文件
            temp_file_path = os.path.join(self.temp_dir, Path(file_path).stem + ".txt")
            with open(temp_file_path, 'w', encoding='utf-8') as f:
                f.write(text_content)
            
            return {
                "content": text_content,
                "paragraphs": paragraphs,
                "temp_file": temp_file_path,
                "file_type": "pdf",
                "file_name": Path(file_path).name
            }
        except ImportError:
            logging.error("缺少PyPDF2库，请安装: pip install PyPDF2")
            raise
        except Exception as e:
            logging.error(f"解析PDF文档失败: {str(e)}")
            raise ValueError(f"解析PDF文档失败: {str(e)}")
    
    def cleanup(self) -> None:
        """清理临时文件"""
        import shutil
        if os.path.exists(self.temp_dir):
            shutil.rmtree(self.temp_dir)
            logging.info(f"清理临时目录: {self.temp_dir}")